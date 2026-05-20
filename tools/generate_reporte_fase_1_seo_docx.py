from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Reporte_Fase_1_SEO.docx"

ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(92, 99, 112)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F5F7"
BORDER = "D9DEE7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_tbl_width(table, width_dxa=9360):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_paragraph_border(paragraph, color=BORDER):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    set_tbl_width(tbl, sum(widths))
    set_table_borders(tbl)
    hdr = tbl.rows[0]
    set_repeat_table_header(hdr)
    for i, label in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
    for row in rows:
        new_row = tbl.add_row()
        set_row_cant_split(new_row)
        cells = new_row.cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if widths[i] <= 1400 else WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(str(value))
    doc.add_paragraph()
    return tbl


def add_finding(doc, title, priority, evidence, impact, action, owner="Desarrollo / SEO"):
    add_heading(doc, title, level=3)
    rows = [
        ("Prioridad", priority),
        ("Evidencia", evidence),
        ("Impacto", impact),
        ("Acción recomendada", action),
        ("Responsable sugerido", owner),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    set_tbl_width(tbl, 9360)
    set_table_borders(tbl)
    for label, value in rows:
        new_row = tbl.add_row()
        set_row_cant_split(new_row)
        cells = new_row.cells
        widths = [1900, 7460]
        for i, cell in enumerate(cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell, top=100, bottom=100)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cells[0], LIGHT_GRAY)
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(value)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 22, ACCENT),
        ("Subtitle", 12, MUTED),
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 14, RGBColor(32, 43, 58)),
        ("Heading 3", 12, RGBColor(32, 43, 58)),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        if "Heading" in style_name or style_name == "Title":
            style.font.bold = True

    for list_style in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[list_style]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)

    header = section.header
    p = header.paragraphs[0]
    p.text = "Reporte Fase 1 SEO - Solo Empleos"
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = MUTED
    set_paragraph_border(p)

    footer = section.footer
    p = footer.paragraphs[0]
    p.add_run("Página ")
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = MUTED
    add_page_number(p)


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Reporte Fase 1 SEO")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Solo Empleos | Auditoría técnica inicial del sitio productivo")
    doc.add_paragraph()

    meta_rows = [
        ("Sitio auditado", "https://soloempleos.click/"),
        ("Páginas revisadas", "/, /gdl/inicio/, /mty/inicio/, /gdl/contacto/, /mty/contacto/"),
        ("Entorno local revisado", "http://localhost:3011/"),
        ("Fecha de reporte", "17 de mayo de 2026"),
        ("Enfoque", "SEO técnico básico, indexabilidad, estructura HTML, recursos públicos, rendimiento móvil y evidencia Lighthouse."),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    set_tbl_width(tbl, 9360)
    set_table_borders(tbl)
    for label, value in meta_rows:
        new_row = tbl.add_row()
        set_row_cant_split(new_row)
        cells = new_row.cells
        for i, width in enumerate([2200, 7160]):
            set_cell_width(cells[i], width)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cells[0], LIGHT_BLUE)
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(value)

    doc.add_page_break()

    add_heading(doc, "1. Resumen ejecutivo", level=1)
    add_body(
        doc,
        "El sitio productivo se encuentra disponible y las páginas públicas principales responden correctamente con HTTP 200. "
        "También se confirmó que las páginas incluyen idioma `es` y viewport móvil, lo que establece una base mínima para rastreo y experiencia móvil.",
    )
    add_body(
        doc,
        "La fase 1 detectó oportunidades claras y accionables en SEO técnico. Las más importantes son la ausencia de meta descriptions, "
        "canonical tags, encabezados H1, robots.txt, sitemap.xml y favicon. Además, la página de Guadalajara presenta un problema de rendimiento "
        "relevante por imágenes pesadas, con LCP de 10.8 segundos en Lighthouse.",
    )
    add_body(
        doc,
        "Conclusión operativa: el sitio está activo, pero necesita una normalización técnica básica antes de avanzar a optimización de contenido, "
        "indexación fina o mejoras de conversión. Las correcciones recomendadas son de bajo riesgo y alto impacto para claridad semántica, rastreo y experiencia móvil.",
    )

    table(
        doc,
        ["Área", "Estado actual", "Riesgo", "Siguiente acción"],
        [
            ("Indexabilidad", "Sin robots.txt ni sitemap.xml", "Alto", "Crear ambos archivos y referenciar sitemap en robots.txt."),
            ("Metadatos", "Sin meta descriptions ni canonicals", "Alto", "Agregar valores únicos por página."),
            ("Semántica HTML", "H1 ausente en páginas principales", "Alto", "Convertir titulares principales a H1 y ordenar H2/H3."),
            ("Rendimiento", "LCP alto en /gdl/inicio/", "Alto", "Optimizar imágenes de vacantes y entrega de recursos."),
            ("Acabado técnico", "Favicon y recurso WhatsApp con 404", "Medio", "Corregir archivos públicos y rutas de despliegue."),
        ],
        [1700, 3000, 1200, 3460],
    )

    add_heading(doc, "2. Alcance revisado", level=1)
    add_body(doc, "La revisión cubrió el sitio productivo y una comprobación del proyecto local para contrastar estructura, recursos y comportamiento visual.")
    for item in [
        "Sitio productivo: https://soloempleos.click/",
        "Páginas públicas principales: home, Guadalajara, Monterrey y páginas de contacto por ciudad.",
        "Archivos de rastreo esperados: robots.txt y sitemap.xml.",
        "Estructura HTML: títulos, meta descriptions, canonicals, encabezados y recursos estáticos.",
        "Validación móvil visual con Playwright y mediciones Lighthouse disponibles.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Hallazgos prioritarios", level=1)
    add_heading(doc, "3.1 Prioridad alta", level=2)
    add_finding(
        doc,
        "Faltan meta descriptions en páginas principales",
        "Alta",
        "Lighthouse reporta: Document does not have a meta description. Afecta a /, /gdl/inicio/, /mty/inicio/, /gdl/contacto/ y /mty/contacto/.",
        "Google tiene menos contexto para interpretar cada página y puede generar snippets menos controlados o menos atractivos.",
        "Agregar una meta description única por URL, con ciudad e intención de búsqueda cuando aplique.",
    )
    add_finding(
        doc,
        "No existe robots.txt",
        "Alta",
        "https://soloempleos.click/robots.txt y http://localhost:3011/robots.txt devuelven 404.",
        "No hay instrucciones básicas para rastreadores ni referencia explícita al sitemap.",
        "Crear robots.txt, permitir páginas públicas, declarar sitemap y definir tratamiento de /admin/.",
    )
    add_finding(
        doc,
        "No existe sitemap.xml",
        "Alta",
        "https://soloempleos.click/sitemap.xml y http://localhost:3011/sitemap.xml devuelven 404.",
        "Se dificulta declarar a buscadores las URLs importantes del sitio y sus variantes canónicas.",
        "Crear sitemap.xml con las páginas públicas principales y actualizarlo cuando se agreguen nuevas rutas relevantes.",
    )
    add_finding(
        doc,
        "No hay encabezados H1",
        "Alta",
        "Todas las páginas principales revisadas tienen H1 = 0. Inicio usa H2 para el texto principal y contacto usa H3.",
        "La página pierde claridad semántica sobre su tema principal, lo que afecta SEO y accesibilidad.",
        "Definir un H1 único por página y reordenar la jerarquía de subtítulos.",
    )
    add_finding(
        doc,
        "Rendimiento bajo en /gdl/inicio/",
        "Alta",
        "Lighthouse: Performance 61, LCP 10.8 s, peso total estimado 19,523 KiB y ahorro potencial de imágenes de 16,572 KiB.",
        "La experiencia móvil es lenta y puede afectar Core Web Vitals, retención de usuarios y visibilidad orgánica.",
        "Comprimir imágenes, servir formatos modernos, definir tamaños responsivos y revisar lazy-loading/priority del contenido crítico.",
    )
    add_finding(
        doc,
        "Falta favicon",
        "Alta",
        "Lighthouse reporta 404 en https://soloempleos.click/favicon.ico.",
        "Indica acabado técnico incompleto y genera error de consola o solicitudes fallidas.",
        "Agregar favicon.ico y, si es posible, set de iconos modernos para navegador y dispositivos.",
    )

    doc.add_page_break()
    add_heading(doc, "3.2 Prioridad media", level=2)
    for title, evidence, impact, action in [
        (
            "No hay etiquetas canonical",
            "Todas las páginas revisadas reportan canonical faltante.",
            "Hay menos control sobre URL canónica, especialmente con variantes con o sin slash.",
            "Agregar canonical absoluto y consistente por página.",
        ),
        (
            "/admin/ es público y responde 200",
            "https://soloempleos.click/admin/ devuelve 200.",
            "Si no debe indexarse, puede quedar expuesto en resultados o rastreos.",
            "Definir política: bloquear en robots.txt, agregar noindex o proteger acceso según objetivo real.",
        ),
        (
            "Recurso de WhatsApp no encontrado en productivo",
            "https://soloempleos.click/shared/img/whatsapp.svg devuelve 404; en el repo local sí existe pages/shared/img/whatsapp.svg.",
            "Puede generar iconos rotos en tarjetas con WhatsApp y errores técnicos visibles.",
            "Corregir ruta pública o proceso de despliegue del recurso compartido.",
        ),
        (
            "Imágenes grandes en uploads",
            "Hay imágenes de vacantes mayores a 1 MB y hasta aproximadamente 2 MB; videos hero pesan 2.3 MB a 3.0 MB.",
            "Aumenta el tiempo de carga, especialmente en móvil.",
            "Crear flujo de optimización y límites de peso para imágenes subidas.",
        ),
        (
            "Orden de encabezados mejorable",
            "Lighthouse reporta orden no secuencial de encabezados en /gdl/inicio/.",
            "Afecta accesibilidad, navegación por lectores de pantalla y claridad semántica.",
            "Revisar jerarquía después de introducir H1 y normalizar H2/H3.",
        ),
    ]:
        add_finding(doc, title, "Media", evidence, impact, action)

    add_heading(doc, "3.3 Prioridad baja", level=2)
    for item in [
        "Los títulos SEO existen, pero son genéricos. Conviene incluir intención de búsqueda, vacantes/empleos y ciudad.",
        "El alt text está presente, pero usa valores genéricos como `Vacante`. En fase 2 o 3 conviene generar textos descriptivos cuando exista información suficiente.",
        "El enlace `Regresar a página anterior` usa `javascript:history.go(-1)`. No es crítico para SEO, pero es frágil para usuarios que aterrizan desde buscador.",
        "La búsqueda `site:soloempleos.click` no arrojó resultados en la herramienta consultada. Esto debe confirmarse en Google Search Console antes de concluir indexación real.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. Evidencia técnica", level=1)
    add_heading(doc, "4.1 HTTP productivo", level=2)
    table(
        doc,
        ["URL", "Estado", "Observación"],
        [
            ("https://soloempleos.click/", "200", "Home disponible"),
            ("https://soloempleos.click/gdl/inicio/", "200", "Guadalajara disponible"),
            ("https://soloempleos.click/mty/inicio/", "200", "Monterrey disponible"),
            ("https://soloempleos.click/gdl/contacto/", "200", "Contacto Guadalajara disponible"),
            ("https://soloempleos.click/mty/contacto/", "200", "Contacto Monterrey disponible"),
            ("https://soloempleos.click/robots.txt", "404", "Falta archivo"),
            ("https://soloempleos.click/sitemap.xml", "404", "Falta archivo"),
        ],
        [5200, 1000, 3160],
    )

    add_heading(doc, "4.2 Metadatos productivos", level=2)
    table(
        doc,
        ["URL", "Meta description", "Canonical", "H1"],
        [
            ("/", "Falta", "Falta", "0"),
            ("/gdl/inicio/", "Falta", "Falta", "0"),
            ("/mty/inicio/", "Falta", "Falta", "0"),
            ("/gdl/contacto/", "Falta", "Falta", "0"),
            ("/mty/contacto/", "Falta", "Falta", "0"),
        ],
        [3300, 2300, 2300, 1460],
    )

    add_heading(doc, "4.3 Lighthouse", level=2)
    table(
        doc,
        ["Página", "Performance", "SEO", "Accesibilidad", "Buenas prácticas", "LCP", "CLS", "TBT"],
        [
            ("/", "91", "91", "96", "96", "3.1 s", "0.099", "0 ms"),
            ("/gdl/inicio/", "61", "91", "93", "96", "10.8 s", "0.053", "500 ms"),
        ],
        [1900, 1200, 850, 1300, 1500, 1100, 760, 750],
    )

    add_heading(doc, "5. Plan recomendado para fase 2", level=1)
    add_body(
        doc,
        "La fase 2 debe cerrar primero los puntos de rastreo, semántica y metadatos. Una vez hecho esto, el sitio queda mejor preparado para optimización de contenido, medición en Search Console y mejoras de rendimiento profundas.",
    )
    for step in [
        "Crear robots.txt y sitemap.xml.",
        "Agregar meta descriptions y canonicals por página.",
        "Convertir encabezados principales a H1 y ordenar jerarquía H2/H3.",
        "Agregar favicon.",
        "Corregir despliegue del recurso shared/img/whatsapp.svg.",
        "Definir si /admin/ debe bloquearse o marcarse como noindex.",
        "Preparar optimización de imágenes para fase 3, priorizando vacantes de /gdl/inicio/.",
    ]:
        add_number(doc, step)

    doc.add_page_break()
    add_heading(doc, "6. Criterios de aceptación sugeridos", level=1)
    for item in [
        "robots.txt y sitemap.xml responden 200 en productivo.",
        "Cada página pública principal tiene title, meta description, canonical y un H1 único.",
        "Lighthouse deja de reportar meta description faltante y encabezados principales ausentes.",
        "favicon.ico y shared/img/whatsapp.svg responden 200.",
        "La política de /admin/ queda documentada e implementada.",
        "Se establece una línea base de rendimiento para comparar después de optimizar imágenes.",
    ]:
        add_bullet(doc, item)

    doc.core_properties.title = "Reporte Fase 1 SEO - Solo Empleos"
    doc.core_properties.subject = "Auditoría técnica SEO inicial"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "SEO, Solo Empleos, auditoría, Lighthouse, robots.txt, sitemap.xml"
    doc.core_properties.comments = "Generado a partir de Reporte_Fase_1_SEO.md"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
