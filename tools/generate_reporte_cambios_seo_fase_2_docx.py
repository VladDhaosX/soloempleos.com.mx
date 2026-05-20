from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Reporte_Cambios_SEO_Fase_2.docx"

ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(91, 101, 115)
BORDER = "D9E2F3"
HEADER_FILL = "EAF2F8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_run(run, size=11, bold=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        style_run(r, bold=True)
        rest = text[len(bold_prefix):]
        if rest:
            r = p.add_run(rest)
            style_run(r)
    else:
        r = p.add_run(text)
        style_run(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        style_run(r)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, HEADER_FILL)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        style_run(r, size=9.5, bold=True, color=ACCENT)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.text = ""
            set_cell_width(cell, widths[idx])
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            style_run(r, size=9)
    doc.add_paragraph()
    return table


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 22, ACCENT),
        ("Heading 1", 15, ACCENT),
        ("Heading 2", 12.5, RGBColor(45, 55, 72)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Solo Empleos | Reporte de cambios SEO Fase 2"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.runs[0].font.name = "Arial"
    hp.runs[0].font.size = Pt(8)
    hp.runs[0].font.color.rgb = MUTED

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Documento de seguimiento tecnico"
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.runs[0].font.name = "Arial"
    fp.runs[0].font.size = Pt(8)
    fp.runs[0].font.color.rgb = MUTED
    return doc


def build():
    doc = setup_document()

    title = doc.add_paragraph(style="Title")
    title.add_run("Reporte de Cambios SEO - Fase 2")

    subtitle = doc.add_paragraph()
    r = subtitle.add_run("Solo Empleos | 2026-05-20 | Cambios aplicados en local, pendientes de despliegue a produccion")
    style_run(r, size=11, color=MUTED)

    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ["Dominio revisado", "https://soloempleos.click/"],
            ["Servidor local de verificacion", "http://localhost:3000/"],
            ["Estado", "Local corregido; produccion aun refleja el estado anterior"],
            ["Objetivo", "Documentar cambios de SEO tecnico basico y medicion antes/despues"],
        ],
        [2300, 7060],
    )

    doc.add_heading("Resumen ejecutivo", level=1)
    add_paragraph(
        doc,
        "Durante la Fase 2 se corrigieron los principales pendientes de SEO tecnico basico detectados en la auditoria inicial: metadatos, encabezados, rastreo, sitemap, favicon, indexabilidad del panel administrativo y alt text menos generico para vacantes.",
    )
    add_paragraph(
        doc,
        "La medicion confirma que produccion todavia esta en estado anterior, mientras que el entorno local ya muestra las correcciones listas para despliegue.",
    )

    doc.add_heading("Alcance", level=1)
    add_bullets(
        doc,
        [
            "Paginas publicas: /, /gdl/inicio/, /mty/inicio/, /gdl/contacto/, /mty/contacto/.",
            "Rutas tecnicas: /robots.txt, /sitemap.xml, /favicon.ico y /admin/.",
            "Comparacion: produccion actual antes del despliegue vs entorno local despues de cambios.",
        ],
    )

    doc.add_heading("Comparacion antes vs despues", level=1)
    add_table(
        doc,
        ["Senal SEO", "Publicado actual", "Local despues"],
        [
            ["Home con meta description", "No", "Si"],
            ["Home con canonical", "No", "Si"],
            ["Home con h1", "0", "1"],
            ["/gdl/inicio/ description/canonical/h1", "No / No / 0", "Si / Si / 1"],
            ["/mty/inicio/ description/canonical/h1", "No / No / 0", "Si / Si / 1"],
            ["/gdl/contacto/ description/canonical/h1", "No / No / 0", "Si / Si / 1"],
            ["/mty/contacto/ description/canonical/h1", "No / No / 0", "Si / Si / 1"],
            ["/robots.txt", "404", "200"],
            ["/sitemap.xml", "404", "200"],
            ["/favicon.ico", "404", "200"],
            ["/admin/ noindex", "No", "Si"],
        ],
        [4700, 2300, 2360],
    )

    doc.add_heading("Cambios realizados", level=1)
    changes = [
        ("Metadatos y encabezados", "Titulos SEO, meta descriptions, canonicals absolutos y un h1 unico en cada pagina publica principal."),
        ("Rastreo", "robots.txt agregado con Allow para paginas publicas, Disallow para /admin/ y referencia al sitemap."),
        ("Sitemap", "sitemap.xml agregado con las cinco URLs publicas principales."),
        ("Favicon", "favicon.ico agregado para eliminar el 404 tecnico."),
        ("Admin", "Meta robots y cabecera X-Robots-Tag noindex, nofollow."),
        ("Semantica visual", "Clase .sr-only agregada para ocultar correctamente el h1 semantico de la landing."),
        ("Alt text", "Vacantes generadas por SSR ahora incluyen ciudad y fecha disponible."),
    ]
    add_table(doc, ["Area", "Cambio"], changes, [2500, 6860])

    doc.add_heading("Archivos afectados", level=1)
    add_bullets(
        doc,
        [
            "HTML publicos: pages/index.html, pages/gdl/inicio/index.html, pages/mty/inicio/index.html, pages/gdl/contacto/index.html, pages/mty/contacto/index.html.",
            "Archivos tecnicos nuevos: pages/robots.txt, pages/sitemap.xml, pages/favicon.ico.",
            "Servidor: server.js.",
            "Admin: admin/index.html.",
            "Estilos compartidos: pages/shared/css/base.css.",
        ],
    )

    doc.add_heading("Verificacion local", level=1)
    add_table(
        doc,
        ["Ruta", "Resultado local"],
        [
            ["/", "200, description presente, canonical presente, h1 = 1"],
            ["/gdl/inicio/", "200, description presente, canonical presente, h1 = 1"],
            ["/mty/inicio/", "200, description presente, canonical presente, h1 = 1"],
            ["/gdl/contacto/", "200, description presente, canonical presente, h1 = 1"],
            ["/mty/contacto/", "200, description presente, canonical presente, h1 = 1"],
            ["/robots.txt", "200"],
            ["/sitemap.xml", "200"],
            ["/favicon.ico", "200"],
            ["/admin/", "200, meta robots y X-Robots-Tag presentes"],
        ],
        [2600, 6760],
    )

    doc.add_heading("Estado publicado antes del despliegue", level=1)
    add_table(
        doc,
        ["Ruta", "Estado publicado actual"],
        [
            ["/", "200, sin description, sin canonical, h1 = 0"],
            ["/gdl/inicio/", "200, sin description, sin canonical, h1 = 0"],
            ["/mty/inicio/", "200, sin description, sin canonical, h1 = 0"],
            ["/gdl/contacto/", "200, sin description, sin canonical, h1 = 0"],
            ["/mty/contacto/", "200, sin description, sin canonical, h1 = 0"],
            ["/robots.txt", "404"],
            ["/sitemap.xml", "404"],
            ["/favicon.ico", "404"],
            ["/admin/", "200, sin meta robots"],
        ],
        [2600, 6760],
    )

    doc.add_heading("Linea base Lighthouse", level=1)
    add_table(
        doc,
        ["Pagina", "Performance", "SEO", "Accesibilidad", "Buenas practicas", "LCP", "CLS", "TBT"],
        [
            ["/", "91", "91", "96", "96", "3.1 s", "0.099", "0 ms"],
            ["/gdl/inicio/", "61", "91", "93", "96", "10.8 s", "0.053", "500 ms"],
        ],
        [1900, 1100, 800, 1300, 1400, 1100, 850, 910],
    )
    add_paragraph(
        doc,
        "Nota: Fase 2 atiende principalmente SEO tecnico. La optimizacion profunda de rendimiento, especialmente imagenes pesadas en /gdl/inicio/, corresponde a Fase 3.",
    )

    doc.add_heading("Proxima medicion despues del despliegue", level=1)
    add_bullets(
        doc,
        [
            "Confirmar en produccion que las cinco paginas publicas tengan description, canonical y un h1 unico.",
            "Confirmar que /robots.txt, /sitemap.xml y /favicon.ico respondan 200.",
            "Confirmar que /admin/ tenga noindex, nofollow.",
            "Ejecutar Lighthouse en home y /gdl/inicio/.",
            "Comparar contra la linea base de Fase 1.",
        ],
    )

    doc.add_heading("Conclusion", level=1)
    add_paragraph(
        doc,
        "La Fase 2 queda implementada en el proyecto local y lista para despliegue. La comparacion antes vs despues muestra que los cambios corrigen los principales problemas de SEO tecnico basico detectados en Fase 1. El paso pendiente es publicar los archivos y repetir la medicion en produccion para documentar el impacto real en el dominio.",
    )

    doc.core_properties.title = "Reporte de Cambios SEO - Fase 2"
    doc.core_properties.subject = "Seguimiento de cambios SEO tecnico basico"
    doc.core_properties.author = "VladApps"
    doc.core_properties.keywords = "SEO, Solo Empleos, fase 2, robots.txt, sitemap.xml, canonical, meta description"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
